"""
CrawlSpider for UPR Websites
"""

# Importing necessary modules for the spider to function
# Spider class allows for web crawling
from scrapy.spiders import Spider
# LinkExtractor helps in extracting links from web pages
from scrapy.linkextractors import LinkExtractor
# ItemLoader helps in loading data into defined item models
from scrapy.loader import ItemLoader
# These are used for defining items, their fields, and making HTTP requests.
from scrapy import Item, Field, Request  
# This is used for data preprocessing during item loading.
from itemloaders.processors import MapCompose
# CrawlerProcess is used to run the spider
from scrapy.crawler import CrawlerProcess
# PyPDF2 is a Python library for reading and manipulating PDF files.
import PyPDF2
# The docx module allows for reading and writing Microsoft Word .docx files.
import docx
# The tempfile module creates temporary files and directories.
import tempfile
# The os module provides a way to interact with the operating system
import os


def join_text(text_list: str) -> str:
    texts = " ".join(text_list)
    return texts


def clean_text(text: str) -> str:
    return text.replace('\n', '').replace('\r', '').replace('\t', '').replace("/", '').strip()


class WebSite(Item):
    url = Field()
    title = Field()
    content = Field(serializer=join_text)


class Document(Item):
    url = Field()
    content = Field()


class UprSpider(Spider):
    """
    Spider to crawl UPR websites.
    """


    def __init__(self):
        if os.path.exists('data.json'):
            os.remove('data.json')

    name = 'upr_spider'  # Unique identifier for the spider
    # Domains that the spider is allowed to crawl
    allowed_domains = ['upr.edu.cu']
    start_urls = [
        'http://www.upr.edu.cu/',
        'https://blogcrai.upr.edu.cu/',
        'https://crai.upr.edu.cu/'
    ]  # URLs to start crawling from

    custom_settings = {
        'FEEDS': {
            'data.json': {
                'format': 'json',
                'encoding': 'utf-8',
                'store_empty': True,
                'indent': 4,
            }
        }
    }  # Custom settings for the spider

    def parse(self, response):
        """
        Parses the response from the website.
        Args:
            response (scrapy.http.Response): The response object containing the web page data.
        Returns:
            scrapy.loader.ItemLoader: An ItemLoader object containing scraped data.
        """
        pdf_links = LinkExtractor(allow=[r".pdf$", r'article/download/']).extract_links(response)
        doc_links = LinkExtractor(allow=('.doc$|.docx$')).extract_links(response)
        txt_links = LinkExtractor(allow=r".txt$").extract_links(response)

        for link in pdf_links:
            yield Request(link.url, callback=self.parse_pdf)

        for link in doc_links:
            yield Request(link.url, callback=self.parse_docx)

        for link in txt_links:
            yield Request(link.url, callback=self.parse_txt)

        item = ItemLoader(
            item=WebSite(), response=response)  # Loading data into WebSite item model
        item.add_value('url', response.url)  # Adding URL to the item
        # Extracting and adding title to the item
        item.add_xpath('title', '/html/head/title/text()',
                       MapCompose(clean_text))
        # Extracting and adding content to the item
        item.add_xpath(
            'content', '//body//text()[not(parent::script) and not(parent::style)]', MapCompose(clean_text))
        yield item.load_item()  # Returning the loaded item

        all_links = LinkExtractor(deny_extensions=[
                    'rar', 'jpg', 'jpeg', 'gif', 'png', 'ppt', 'pptx', 'pdf', 'doc', 'docx', 'txt', 'xls', 'db', 'zip', 'dpt', 'exe', 'mso', 'wmz', 'sav', 'tmp'
                    ]).extract_links(response)

        for link in all_links:
            yield Request(link.url, callback=self.parse)

    def parse_pdf(self, response):
        """
        Parses documents found in the website.
        Args:
            response (scrapy.http.Response): The response object containing the document data.
        Returns:
            scrapy.loader.ItemLoader: An ItemLoader object containing scraped document data.
        """
        print("*"*20+"PDF"+"*"*20)
        item = ItemLoader(item=Document(), response=response)
        url = response.url
        content = ''
        try:
            with tempfile.NamedTemporaryFile(delete=True) as temp_file:  # Creating a temporary file
                temp_file.write(response.body)
                temp_file.seek(0)
                pdf_reader = PyPDF2.PdfReader(temp_file) # Reading temp_file as PDF
                number_pages = len(pdf_reader.pages) # Number of pages of pdf
                for page_num in range(number_pages):
                    content += pdf_reader.pages[page_num].extract_text() # Extract text from pdf page
        except Exception as e:
            print(f'Error:\n {e}')
        item.add_value("url", url)
        item.add_value("content", content)
        return item.load_item()

    def parse_docx(self, response):
        """
        Parses documents in DOCX format found in the website.
        Args:
            response (scrapy.http.Response): The response object containing the document data.
        Returns:
            scrapy.loader.ItemLoader: An ItemLoader object containing scraped document data.
        """
        print("*"*20+"DOC"+"*"*20)
        
        item = ItemLoader(item=Document(), response=response)
        url = response.url
        content = ""
        try:
            with tempfile.NamedTemporaryFile(delete=True) as temp_file:
                temp_file.write(response.body)
                temp_file.seek(0)
                doc = docx.Document(temp_file.name)
                for paragraph in doc.paragraphs:
                    content += paragraph.text
        except Exception as e:
            print(f'Error:\n {e}')
        item.add_value("url", url)
        item.add_value("content", content)
        return item.load_item()

    def parse_txt(self, response):
        """
        Parses documents in TXT format found in the website.

        Args:
            response (scrapy.http.Response): The response object containing the document data.

        Returns:
            scrapy.loader.ItemLoader: An ItemLoader object containing scraped document data.
        """
        print("*"*20+"TXT"+"*"*20)
        item = ItemLoader(item=Document(), response=response)
        url = response.url
        content = response.body.decode('utf-8')
        item.add_value("url", url)
        item.add_value("content", content)
        return item.load_item()


if __name__ == '__main__':
    process = CrawlerProcess()  # Creating a crawler process
    process.crawl(UprSpider)  # Crawling using the UprSpider
    process.start()  # Starting the crawler process
